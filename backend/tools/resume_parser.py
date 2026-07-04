"""
tools/resume_parser.py — parses an uploaded resume (PDF or DOCX) and
produces a structured profile.json using a single Gemini call.

Usage:
    from tools.resume_parser import parse_resume
    profile = parse_resume("path/to/resume.pdf")   # or .docx
    # profile is a dict — also written to data/profile.json
"""

import copy
import json
import logging
from pathlib import Path

import google.generativeai as genai

from config import GEMINI_API_KEY, GEMINI_MODEL, profile_path

log = logging.getLogger(__name__)
genai.configure(api_key=GEMINI_API_KEY)

# ── Default profile template (null = not found in resume) ────────────────────

DEFAULT_PROFILE: dict = {
    "name": None,
    "email": None,
    "phone": None,
    "location": None,

    # Skills — plain list of strings
    "skills": [],

    # Experience in years per technology/domain, e.g. {"python": 3, "aws": 2}
    "experience_years": {},

    # Highest degree + field, e.g. "B.Tech Computer Science"
    "education": None,

    # List of certification names
    "certifications": [],

    # Short descriptions of notable projects (used to judge depth vs. job reqs)
    "projects": [],

    # Job titles the candidate is targeting, e.g. ["backend engineer", "ML engineer"]
    "preferred_roles": [],

    # Candidate's expected salary range in USD/year (null if not on resume)
    "salary_min": None,
    "salary_max": None,

    # Location preferences, e.g. ["remote", "San Francisco", "Bangalore"]
    "location_preferences": [],

    # Total years of professional experience (inferred)
    "total_experience_years": None,

    # Top keywords to use when querying job APIs (auto-derived from skills + roles)
    "search_keywords": [],
}


# ── Text extraction ────────────────────────────────────────────────────────────

def _extract_text_pdf(path: str) -> str:
    """Extract text from a PDF resume using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        raise RuntimeError("pdfplumber not installed. Run: pip install pdfplumber")

    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text(x_tolerance=2, y_tolerance=3)
            if page_text:
                text_parts.append(page_text)

    full_text = "\n".join(text_parts)
    if not full_text.strip():
        raise ValueError(
            f"No text extracted from {path}. "
            "The PDF may be scanned/image-based. "
            "Please provide a text-based PDF or convert to DOCX."
        )
    return full_text


def _extract_text_docx(path: str) -> str:
    """Extract text from a DOCX resume using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also pull text from tables (some resumes use tables for layout)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text and cell_text not in paragraphs:
                    paragraphs.append(cell_text)

    return "\n".join(paragraphs)


def extract_text(path: str) -> str:
    """
    Dispatch to the right extractor based on file extension.
    Returns raw resume text.
    """
    suffix = Path(path).suffix.lower()
    if suffix == ".pdf":
        return _extract_text_pdf(path)
    elif suffix in (".docx", ".doc"):
        return _extract_text_docx(path)
    elif suffix in (".txt", ".md"):
        return Path(path).read_text(encoding="utf-8")
    else:
        raise ValueError(f"Unsupported resume format: {suffix}. Use PDF or DOCX.")


# ── Gemini extraction ─────────────────────────────────────────────────────────

_SYSTEM_PROMPT = """\
You are a resume parser and LaTeX typesetter. From the resume text, do two things in one response:
1. Extract structured profile data as JSON.
2. Produce a complete, compilable LaTeX resume document.

Return ONLY a valid JSON object with exactly these keys. No markdown, no backticks, no explanation.

{
  "name": string or null,
  "email": string or null,
  "phone": string or null,
  "location": string or null,
  "skills": [list of skill strings],
  "experience_years": {technology_or_domain: years_as_number},
  "education": string or null,
  "certifications": [list of certification name strings],
  "projects": [
    {
      "name": string,
      "description": string (1-2 sentences summarising what was built and tech used),
      "technologies": [list of tech strings]
    }
  ],
  "preferred_roles": [list of job title strings the candidate is targeting],
  "salary_min": number or null,
  "salary_max": number or null,
  "location_preferences": [list of location preference strings],
  "total_experience_years": number or null,
  "search_keywords": [top 8-12 keywords for job search APIs, derived from skills + roles],
  "latex": string (a complete compilable LaTeX document — see LATEX RULES below)
}

PROFILE EXTRACTION RULES:
- For experience_years, infer years per technology from job dates + project mentions.
  If someone worked 2018-2021 as a Python developer, that is 3 years of Python.
- For projects, include only substantial ones (not trivial examples).
- search_keywords must be specific and job-searchable: "python", "machine learning",
  "backend engineer" — not generic words like "software" or "developer".
- Salary values should be annual USD. If the resume mentions a different currency,
  convert approximately. If no salary is mentioned, use null.
- Do not use double quotes inside string values — use single quotes or rephrase instead.
  e.g. write Grad-CAM not "Grad-CAM", write CNN not "CNN".
- For skills: include skills explicitly mentioned AND obvious ones clearly demonstrated
  by projects or job responsibilities (e.g. if they built REST APIs, include REST API).

LATEX RULES:
- The "latex" value must be a single JSON string containing a complete, compilable LaTeX document.
- Use ONLY these packages (no others — the compile service may not have additional ones):
  geometry, titlesec, enumitem, hyperref, parskip, fontenc, inputenc
- Do NOT use fontawesome, fontawesome5, marvosym, pifont, or any icon/symbol package.
- Set hyperref to monochrome: colorlinks=false, so links have no colour.
- Template structure — use exactly this order:
  \\documentclass[11pt,a4paper]{article}
  (package declarations, geometry with 0.8in margins, titlesec formatting, parskip)
  \\begin{document}
  Header: candidate name in Large+bold, then a single contact line using plain text:
    Email | Phone | Location  (and LinkedIn/GitHub if present in the resume)
    Use plain text labels — NO icon packages.
  Sections in this order (omit any section with no content):
    Summary (only if the resume has one), Education, Technical Skills,
    Experience, Projects, Certifications/Achievements
  \\end{document}
- Use \\noindent\\rule{\\linewidth}{0.4pt} after each section heading as a divider.
- Style: clean, single-column, monochrome, ATS-friendly. No colours, no tables for layout.
- Content: include ONLY information present in the resume. No placeholders, no invented entries.
- Dates, job titles, company names, project names must match the resume exactly.
"""

_USER_TEMPLATE = """\
Resume text:
---
{resume_text}
---

Return the JSON object now.
"""


def _repair_json(raw: str) -> dict:
    """
    Attempt to fix common Gemini JSON issues before giving up.
    Tries json-repair library first, then truncates at last valid closing brace.
    """
    # Try json-repair if available
    try:
        from json_repair import repair_json
        repaired = repair_json(raw)
        return json.loads(repaired)
    except Exception:
        pass

    # Fallback: truncate at last } and re-parse
    last_brace = raw.rfind("}")
    if last_brace != -1:
        truncated = raw[:last_brace + 1]
        try:
            return json.loads(truncated)
        except json.JSONDecodeError:
            pass

    raise ValueError("Could not repair JSON")


def _call_gemini(resume_text: str) -> tuple[dict, str]:
    """
    Send resume text to Gemini. Returns (profile_dict, latex_str).
    latex_str is empty string if Gemini omitted it or it failed to parse —
    the caller must treat a missing latex as non-fatal.
    """
    model = genai.GenerativeModel(
        model_name=GEMINI_MODEL,
        system_instruction=_SYSTEM_PROMPT,
        generation_config={"response_mime_type": "application/json"},
    )

    # Truncate resume text if extremely long (>8000 chars) to stay within token limits.
    # LaTeX is generated from the same truncated text — still covers all real resumes.
    truncated = resume_text[:8000]
    if len(resume_text) > 8000:
        log.warning("Resume text truncated from %d to 8000 chars for Gemini.", len(resume_text))

    response = model.generate_content(_USER_TEMPLATE.format(resume_text=truncated))
    raw = response.text.strip()

    try:
        data = json.loads(raw)
    except json.JSONDecodeError as e:
        log.warning("Gemini JSON malformed (%s) — attempting repair...", e)
        try:
            data = _repair_json(raw)
        except ValueError:
            log.error("Could not repair Gemini JSON:\n%s", raw)
            raise ValueError(f"Gemini did not return valid JSON: {e}\nRaw output:\n{raw[:500]}")

    # Extract latex separately so a missing/malformed latex never raises here.
    latex = ""
    raw_latex = data.pop("latex", None)
    if isinstance(raw_latex, str) and raw_latex.strip():
        latex = raw_latex.strip()
    elif raw_latex is not None:
        log.warning("Gemini returned a non-string or empty 'latex' field — skipping.")

    return data, latex


# ── Merge with default template ───────────────────────────────────────────────

def _merge_with_defaults(gemini_profile: dict) -> dict:
    """
    Start from DEFAULT_PROFILE so all keys are always present,
    then overwrite with whatever Gemini actually found.
    Lists and dicts are replaced, not merged.
    """
    profile = copy.deepcopy(DEFAULT_PROFILE)
    for key, value in gemini_profile.items():
        if key in profile:
            # Only overwrite if the value is not None (empty lists/dicts are valid)
            if value is not None:
                profile[key] = value
    return profile


# ── Public API ─────────────────────────────────────────────────────────────────

def _store_latex_in_supabase(user_id: str, latex: str) -> None:
    """
    Upsert the base LaTeX resume into profiles.resume_latex.
    Deliberately swallows all exceptions — LaTeX storage is additive and
    must never break the upload pipeline if Supabase is unavailable.
    """
    try:
        from supabase import create_client
        import os
        url = os.getenv("SUPABASE_URL", "")
        key = os.getenv("SUPABASE_SERVICE_KEY", "")
        if not url or not key:
            log.warning("Supabase env vars not set — skipping LaTeX storage.")
            return
        sb = create_client(url, key)
        sb.table("profiles").upsert(
            {"user_id": user_id, "resume_latex": latex},
            on_conflict="user_id",
        ).execute()
        log.info("Stored base LaTeX for user=%s (%d chars)", user_id, len(latex))
    except Exception as e:
        log.warning("Could not store LaTeX in Supabase for user=%s: %s", user_id, e)


def parse_resume(resume_path: str, user_id: str) -> dict:
    """
    Parse a resume file and return a structured profile dict.

    Steps:
      1. Extract plain text from PDF or DOCX.
      2. Send to Gemini — ONE call that returns both profile JSON and base LaTeX.
      3. Merge profile with the default template (ensures all keys exist).
      4. Save profile to data/profile.json.
      5. Store base LaTeX in Supabase profiles.resume_latex (non-fatal if it fails).

    Returns:
        Profile dict with all DEFAULT_PROFILE keys guaranteed to exist.
    """
    resume_path = str(resume_path)
    log.info("Parsing resume: %s", resume_path)
    print(f"  Extracting text from {Path(resume_path).name}...")

    # Step 1: Extract text
    raw_text = extract_text(resume_path)
    char_count = len(raw_text)
    print(f"  Extracted {char_count:,} characters of text.")

    if char_count < 100:
        raise ValueError(
            "Extracted text is too short to be a valid resume. "
            "Check that the file is not blank or image-only."
        )

    # Step 2: Gemini parse (profile + LaTeX in one call)
    print("  Sending to Gemini for structured extraction + LaTeX generation (1 API call)...")
    gemini_profile, latex = _call_gemini(raw_text)

    # Step 3: Merge with defaults
    profile = _merge_with_defaults(gemini_profile)

    # Step 4: Save profile
    out_path = profile_path(user_id)
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(profile, f, indent=2, ensure_ascii=False)

    name_str = profile.get("name") or "Unknown"
    skills_count = len(profile.get("skills", []))
    keywords_count = len(profile.get("search_keywords", []))
    print(f"  Profile built for: {name_str}")
    print(f"    Skills found    : {skills_count}")
    print(f"    Search keywords : {keywords_count}")
    print(f"    Saved to        : {out_path}")

    # Step 5: Store LaTeX (non-fatal — never let this break the upload)
    if latex:
        print(f"  Storing base LaTeX ({len(latex):,} chars) in Supabase...")
        _store_latex_in_supabase(user_id, latex)
        print("  LaTeX stored.")
    else:
        log.warning("Gemini returned no LaTeX for user=%s — skipping storage.", user_id)
        print("  Warning: Gemini returned no LaTeX — resume tailoring will be unavailable until next upload.")

    log.info("Profile saved: %s", out_path)
    return profile


# ── CLI helper (run directly to test) ─────────────────────────────────────────

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 3:
        print("Usage: python -m tools.resume_parser path/to/resume.pdf <user_id>")
        sys.exit(1)

    import logging
    logging.basicConfig(level=logging.INFO)

    result = parse_resume(sys.argv[1], user_id=sys.argv[2])
    print("\nExtracted profile:")
    print(json.dumps(result, indent=2))
